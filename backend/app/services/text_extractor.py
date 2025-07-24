import PyPDF2
import pdfplumber
from docx import Document
import markdown
from typing import Dict, Optional
import logging
from pathlib import Path
import io

logger = logging.getLogger(__name__)

class TextExtractor:
    """Extracts text content from various document formats"""
    
    def __init__(self):
        self.extractors = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.doc': self._extract_docx,  # Try docx extractor for .doc files
            '.txt': self._extract_txt,
            '.md': self._extract_markdown,
            '.rtf': self._extract_txt  # Basic RTF support as plain text
        }
    
    async def extract_text(self, file_content: bytes, filename: str, mime_type: str) -> Dict:
        """Extract text from file content"""
        try:
            file_ext = Path(filename).suffix.lower()
            
            if file_ext not in self.extractors:
                return {
                    "success": False,
                    "error": f"Unsupported file extension: {file_ext}"
                }
            
            # Extract text using appropriate extractor
            extractor = self.extractors[file_ext]
            result = await extractor(file_content, filename)
            
            if result["success"]:
                # Add metadata
                text = result["text"]
                result.update({
                    "word_count": len(text.split()),
                    "char_count": len(text),
                    "line_count": len(text.splitlines()),
                    "filename": filename,
                    "file_extension": file_ext,
                    "mime_type": mime_type
                })
            
            return result
            
        except Exception as e:
            logger.error(f"Text extraction error for {filename}: {e}")
            return {
                "success": False,
                "error": f"Extraction failed: {str(e)}"
            }
    
    async def _extract_pdf(self, file_content: bytes, filename: str) -> Dict:
        """Extract text from PDF using pdfplumber (better than PyPDF2)"""
        try:
            text_content = []
            page_count = 0
            
            with io.BytesIO(file_content) as pdf_buffer:
                with pdfplumber.open(pdf_buffer) as pdf:
                    page_count = len(pdf.pages)
                    
                    for page_num, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text_content.append(page_text)
                        except Exception as e:
                            logger.warning(f"Failed to extract page {page_num + 1} from {filename}: {e}")
                            continue
            
            if not text_content:
                # Fallback to PyPDF2
                return await self._extract_pdf_fallback(file_content, filename)
            
            full_text = "\n\n".join(text_content)
            
            return {
                "success": True,
                "text": full_text,
                "page_count": page_count,
                "extractor": "pdfplumber"
            }
            
        except Exception as e:
            logger.warning(f"pdfplumber failed for {filename}: {e}, trying PyPDF2")
            return await self._extract_pdf_fallback(file_content, filename)
    
    async def _extract_pdf_fallback(self, file_content: bytes, filename: str) -> Dict:
        """Fallback PDF extraction using PyPDF2"""
        try:
            text_content = []
            page_count = 0
            
            with io.BytesIO(file_content) as pdf_buffer:
                pdf_reader = PyPDF2.PdfReader(pdf_buffer)
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(page_text)
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num + 1} from {filename}: {e}")
                        continue
            
            full_text = "\n\n".join(text_content)
            
            return {
                "success": True,
                "text": full_text,
                "page_count": page_count,
                "extractor": "PyPDF2"
            }
            
        except Exception as e:
            logger.error(f"PDF extraction failed for {filename}: {e}")
            return {
                "success": False,
                "error": f"PDF extraction failed: {str(e)}"
            }
    
    async def _extract_docx(self, file_content: bytes, filename: str) -> Dict:
        """Extract text from DOCX files"""
        try:
            with io.BytesIO(file_content) as docx_buffer:
                doc = Document(docx_buffer)
                
                paragraphs = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text)
                
                # Extract text from tables
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            tables_text.append(" | ".join(row_text))
                
                # Combine paragraphs and tables
                all_text = paragraphs + tables_text
                full_text = "\n\n".join(all_text)
                
                return {
                    "success": True,
                    "text": full_text,
                    "paragraph_count": len(paragraphs),
                    "table_count": len(doc.tables),
                    "extractor": "python-docx"
                }
                
        except Exception as e:
            logger.error(f"DOCX extraction failed for {filename}: {e}")
            return {
                "success": False,
                "error": f"DOCX extraction failed: {str(e)}"
            }
    
    async def _extract_txt(self, file_content: bytes, filename: str) -> Dict:
        """Extract text from plain text files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return {
                        "success": True,
                        "text": text,
                        "encoding": encoding,
                        "extractor": "text"
                    }
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            text = file_content.decode('utf-8', errors='replace')
            return {
                "success": True,
                "text": text,
                "encoding": "utf-8 (with errors replaced)",
                "extractor": "text"
            }
            
        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {e}")
            return {
                "success": False,
                "error": f"Text extraction failed: {str(e)}"
            }
    
    async def _extract_markdown(self, file_content: bytes, filename: str) -> Dict:
        """Extract text from Markdown files"""
        try:
            # First decode the markdown
            text_result = await self._extract_txt(file_content, filename)
            if not text_result["success"]:
                return text_result
            
            markdown_text = text_result["text"]
            
            # Convert markdown to plain text (remove formatting)
            html = markdown.markdown(markdown_text)
            
            # Simple HTML tag removal (basic approach)
            import re
            plain_text = re.sub(r'<[^>]+>', '', html)
            plain_text = re.sub(r'\s+', ' ', plain_text).strip()
            
            return {
                "success": True,
                "text": plain_text,
                "original_markdown": markdown_text,
                "encoding": text_result.get("encoding", "utf-8"),
                "extractor": "markdown"
            }
            
        except Exception as e:
            logger.error(f"Markdown extraction failed for {filename}: {e}")
            return {
                "success": False,
                "error": f"Markdown extraction failed: {str(e)}"
            }

# Global instance
text_extractor = TextExtractor()
