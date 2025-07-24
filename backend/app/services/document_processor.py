"""
Document processing service
"""
import logging
from typing import List, Dict, Any
import PyPDF2
import docx
import io
from datetime import datetime

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing different document types"""
    
    def __init__(self):
        self.supported_types = ['pdf', 'docx', 'txt']
    
    async def extract_text(self, content: bytes, file_type: str) -> str:
        """
        Extract text content from document based on file type
        """
        try:
            if file_type == 'pdf':
                return await self._extract_pdf_text(content)
            elif file_type == 'docx':
                return await self._extract_docx_text(content)
            elif file_type == 'txt':
                return await self._extract_txt_text(content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Text extraction failed for {file_type}: {e}")
            raise
    
    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"[Page {page_num + 1}]\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            if not text_content:
                raise ValueError("No text content found in PDF")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            raise ValueError(f"Failed to extract text from PDF: {e}")
    
    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(content)
            doc = docx.Document(docx_file)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            if not text_content:
                raise ValueError("No text content found in DOCX")
            
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {e}")
    
    async def _extract_txt_text(self, content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    if text.strip():
                        return text
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file with any supported encoding")
            
        except Exception as e:
            logger.error(f"TXT text extraction failed: {e}")
            raise ValueError(f"Failed to extract text from TXT: {e}")
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Split text into overlapping chunks for processing
        """
        try:
            if not text or not text.strip():
                return []
            
            # Split by sentences first (simple approach)
            sentences = text.replace('\n', ' ').split('. ')
            
            chunks = []
            current_chunk = ""
            
            for sentence in sentences:
                # Add sentence to current chunk
                test_chunk = current_chunk + sentence + ". "
                
                if len(test_chunk) <= chunk_size:
                    current_chunk = test_chunk
                else:
                    # Current chunk is full, save it and start new one
                    if current_chunk.strip():
                        chunks.append(current_chunk.strip())
                    
                    # Start new chunk with overlap
                    if overlap > 0 and chunks:
                        # Take last part of previous chunk for overlap
                        overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                        current_chunk = overlap_text + sentence + ". "
                    else:
                        current_chunk = sentence + ". "
            
            # Add the last chunk
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            
            return chunks
            
        except Exception as e:
            logger.error(f"Text chunking failed: {e}")
            return [text]  # Return original text as single chunk if chunking fails
    
    def extract_metadata(self, text: str, filename: str) -> Dict[str, Any]:
        """
        Extract additional metadata from document content
        """
        try:
            metadata = {
                "word_count": len(text.split()),
                "character_count": len(text),
                "line_count": len(text.split('\n')),
                "estimated_reading_time": max(1, len(text.split()) // 200),  # ~200 WPM
                "language": self._detect_language(text),
                "has_tables": "| " in text,  # Simple table detection
                "has_numbers": any(char.isdigit() for char in text),
                "processed_at": datetime.utcnow().isoformat()
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Metadata extraction failed: {e}")
            return {}
    
    def _detect_language(self, text: str) -> str:
        """
        Simple language detection (placeholder implementation)
        In production, use a proper language detection library
        """
        try:
            # Very basic language detection based on common words
            english_words = ['the', 'and', 'is', 'in', 'to', 'of', 'a', 'that', 'it', 'with']
            text_lower = text.lower()
            
            english_count = sum(1 for word in english_words if word in text_lower)
            
            if english_count >= 3:
                return "en"
            else:
                return "unknown"
                
        except Exception:
            return "unknown"
    
    def validate_document(self, content: bytes, filename: str) -> Dict[str, Any]:
        """
        Validate document before processing
        """
        validation_result = {
            "valid": True,
            "errors": [],
            "warnings": []
        }
        
        try:
            # Check file size
            if len(content) == 0:
                validation_result["valid"] = False
                validation_result["errors"].append("File is empty")
            
            # Check filename
            if not filename or '.' not in filename:
                validation_result["valid"] = False
                validation_result["errors"].append("Invalid filename")
            
            # Check file extension
            file_extension = filename.split('.')[-1].lower()
            if file_extension not in self.supported_types:
                validation_result["valid"] = False
                validation_result["errors"].append(f"Unsupported file type: {file_extension}")
            
            # Add warnings for large files
            if len(content) > 5 * 1024 * 1024:  # 5MB
                validation_result["warnings"].append("Large file may take longer to process")
            
            return validation_result
            
        except Exception as e:
            logger.error(f"Document validation failed: {e}")
            validation_result["valid"] = False
            validation_result["errors"].append("Validation failed")
            return validation_result
